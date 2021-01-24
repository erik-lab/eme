
# copyright - NOT - copied from online forum
# https://stackoverflow.com/questions/52719258/docx-file-to-text-file-conversion-using-python

"""
extracts the text paragraphs from a word Doc.
"""



from __future__ import print_function

from apiclient import discovery
from httplib2 import Http
from oauth2client import client
from oauth2client import file
from oauth2client import tools
import lxml
from docx import Document


SCOPES = 'https://www.googleapis.com/auth/documents.readonly'
DISCOVERY_DOC = 'https://docs.googleapis.com/$discovery/rest?version=v1'
DOCUMENT_ID = '19OnmdmDOjiO19hu8pJlYTmFnrsH5Q073'


def get_credentials():
    """Gets valid user credentials from storage.

    If nothing has been stored, or if the stored credentials are invalid,
    the OAuth 2.0 flow is completed to obtain the new credentials.

    Returns:
        Credentials, the obtained credential.
    """
    store = file.Storage('token.json')
    credentials = store.get()

    if not credentials or credentials.invalid:
        flow = client.flow_from_clientsecrets('credentials.json', SCOPES)
        credentials = tools.run_flow(flow, store)
    return credentials

def read_doc_element(element):
    """Returns the text in the given ParagraphElement.

        Args:
            element: a ParagraphElement from a Google Doc.
    """
    text_run = element.get('textRun')
    if not text_run:
        return ''
    return text_run.get('content')


def read_doc_elements(elements):
    """Recurses through a list of Structural Elements to read a document's text where text may be
        in nested elements.

        Args:
            elements: a list of Structural Elements.
    """
    text = ''
    for value in elements:
        if 'paragraph' in value:
            elements = value.get('paragraph').get('elements')
            for elem in elements:
                text += read_doc_element(elem)
        elif 'table' in value:
            # The text in table cells are in nested Structural Elements and tables may be
            # nested.
            table = value.get('table')
            for row in table.get('tableRows'):
                cells = row.get('tableCells')
                for cell in cells:
                    text += read_doc_elements(cell.get('content'))
        elif 'tableOfContents' in value:
            # The text in the TOC is also in a Structural Element.
            toc = value.get('tableOfContents')
            text += read_doc_elements(toc.get('content'))
    return text

def read_docx_text(docname):
    fileExtension=docname.split(".")[-1]
    if fileExtension == "docx":
        print(docname)
        document = Document(docname)
        thetext = ""
        for para in document.paragraphs:
            thetext = thetext + unicode(para.text)
    return thetext

def main():
    """Uses the Docs API to print out the text of a document."""
    credentials = get_credentials()
    http = credentials.authorize(Http())
    docs_service = discovery.build(
        'docs', 'v1', http=http, discoveryServiceUrl=DISCOVERY_DOC)
    #  doc = docs_service.documents().get(documentId=DOCUMENT_ID).execute()
    download_file(docs_service, DOCUMENT_ID, item['application/vnd.openxmlformats-officedocument.wordprocessingml.document'])


if __name__ == '__main__':
    main()

"""
****
from docx import Document
import io
import shutil
import os

def convertDocxToText(path):
    for d in os.listdir(path):
        fileExtension=d.split(".")[-1]
        if fileExtension =="docx":
            docxFilename = path + d
            print(docxFilename)
            document = Document(docxFilename)
            textFilename = path + d.split(".")[0] + ".txt"
            with io.open(textFilename,"w", encoding="utf-8") as textFile:
                for para in document.paragraphs:
                    textFile.write(unicode(para.text))

path= "/home/python/resumes/"
convertDocxToText(path)
"""
