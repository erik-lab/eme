from __future__ import print_function
import pickle
import os.path
import io
import json
from apiclient import discovery
from apiclient import errors
from apiclient.http import MediaIoBaseDownload
from httplib2 import Http
from oauth2client import client
from oauth2client import file
from oauth2client import tools
from googleapiclient.discovery import build
from google_auth_oauthlib.flow import InstalledAppFlow
from google.auth.transport.requests import Request
from operator import itemgetter, attrgetter
from monkeylearn import MonkeyLearn
import lxml
from docx import Document

# If modifying these scopes, delete the file drive.pickle.
SCOPES = ['https://www.googleapis.com/auth/drive.metadata.readonly',
          'https://www.googleapis.com/auth/drive.readonly',
          'https://www.googleapis.com/auth/documents.readonly']
DISCOVERY_DOC = 'https://docs.googleapis.com/$discovery/rest?version=v1'
DOCUMENT_ID = '1ht6PMhI5JIcaQLqgsMBQhwOIIlHYr455'

def get_credentials():
    """Gets valid user credentials from storage.

    If nothing has been stored, or if the stored credentials are invalid,
    the OAuth 2.0 flow is completed to obtain the new credentials.

    Returns:
        Credentials, the obtained credential.
    """
    store = file.Storage('token.json')
    credentials = store.get()

    if not credentials or credentials.invalid:
        flow = client.flow_from_clientsecrets('credentials.json', SCOPES)
        credentials = tools.run_flow(flow, store)
    return credentials


def retrieve_all_files(service):
    """Retrieve a list of File resources.

    Args:
      service: Drive API service instance.
    Returns:
      List of File resources.
    """
    result = []
    page_token = None
    while True:
        try:
            param = {}
            if page_token:
                param['pageToken'] = page_token
            param['corpora'] = 'user'
            # param['driveId'] = 'root'
            # param['shared'] = 'false'  #  get API error when using this or "trashed"
            param['includeItemsFromAllDrives'] = 'false'
            param['supportsAllDrives'] = 'false'
            param['orderBy'] = 'name'
            param['fields'] = '*'
            param['q'] = "trashed = false"  # "mimeType contains '.folder' and trashed = false"

            files = service.files().list(**param).execute()

            result.extend(files['files'])
            page_token = files.get('nextPageToken')
            if not page_token:
                break
        except errors.HttpError as error:
            print('An error occurred: %s' % error)
            break
    return result


def download_file(service, file_id, mimeType, filename):
    if "google-apps" in mimeType:
        return
    request = service.files().get_media(fileId=file_id)
    fh = io.FileIO(filename, 'wb')  # io.BytesIO()
    downloader = MediaIoBaseDownload(fh, request)
    done = False
    while done is False:
        status, done = downloader.next_chunk()
        print("Download %d%%." % int(status.progress() * 100))
    return  # use with io.BytesIO  fh.getvalue()


def read_docx_text(docname):
    # TODO expand the logic in Docx reader to include tables, headers, and lists
    fileExtension = docname.split(".")[-1]
    thetext = ""
    if fileExtension == "docx":
        document = Document(docname)
        for para in document.paragraphs:
            thetext = thetext + para.text
    return thetext


def read_paragraph_element(element):
    """Returns the text in the given ParagraphElement.

        Args:
            element: a ParagraphElement from a Google Doc.
    """
    text_run = element.get('textRun')
    if not text_run:
        return ''
    return text_run.get('content')


def read_structural_elements(elements):
    """Recurses through a list of Structural Elements to read a document's text where text may be
        in nested elements.

        Args:
            elements: a list of Structural Elements.
    """
    text = ''
    for value in elements:
        if 'paragraph' in value:
            elements = value.get('paragraph').get('elements')
            for elem in elements:
                text += read_paragraph_element(elem)
        elif 'table' in value:
            # The text in table cells are in nested Structural Elements and tables may be
            # nested.
            table = value.get('table')
            for row in table.get('tableRows'):
                cells = row.get('tableCells')
                for cell in cells:
                    text += read_structural_elements(cell.get('content'))
        elif 'tableOfContents' in value:
            # The text in the TOC is also in a Structural Element.
            toc = value.get('tableOfContents')
            text += read_structural_elements(toc.get('content'))
    return text


def main():
    """Shows basic usage of the Drive v3 API.
    Prints the names and ids of the files the user has access to.
    """

    creds = get_credentials()
    service = build('drive', 'v3', credentials=creds)
    http = creds.authorize(Http())
    docs_service = discovery.build(
        'docs', 'v1', http=http, discoveryServiceUrl=DISCOVERY_DOC)

    # Call the Drive v3 API
    # results = service.files().list(
    #     pageSize=10, fields="nextPageToken, files(id, name)").execute()

    f = retrieve_all_files(service)

    for item in f:
        dtype = item['mimeType']  # get the mime type
        try:
            p = item['parents']
        except:
            p = ['NONE']

        print('%s: \t %s \t %s' % (item['id'], item['name'], item['mimeType']))  # item['id']))
        if item['id'] == item['id']:   # DOCUMENT_ID:
            txt = ''            # initialize the doc text holder
            # todo add more generic logic for any doc type
            fileExtension = item['name'].split(".")[-1]
            if fileExtension == "txt" or fileExtension == 'csv':
                # print('do a txt file')
                download_file(service, item['id'], item['mimeType'], 'temp.txt')
                tf = open('temp.txt', 'r')
                txt = tf.read()
                tf.close()

            elif fileExtension == 'docx':
                # print('do a docx')
                download_file(service, item['id'], item['mimeType'], item['name'])
                txt = read_docx_text(item['name'])

            elif 'google-apps.document' in item['mimeType']:
                # print('Do a google Doc')
                doc = docs_service.documents().get(documentId=item['id']).execute()
                doc_content = doc.get('body').get('content')
                txt = read_structural_elements(doc_content)

            elif 'google-apps.spreadsheet' in item['mimeType']:  # todo add google sheet logic is same for docs and presentations?
                print('Do google Sheet')

            else:
                print('unknown Doc Type')

            # at this point we have the text from the document
            # now extract keywords from doc name and content and
            # append all to a keyword list
            # todo create a keyword function for the item that returns a list

            if txt:
                ml = MonkeyLearn('ed48b60fefe026fce0c8220fbbfbdad812c5a7c0')
                data = [txt]
                model_id = 'ex_YCya9nrn'
                result = ml.extractors.extract(model_id, data)
                # json_data = json.loads(result)
                json_data = result.body[0]
                print("%d Keywords" % len(json_data['extractions']))
                # for keyword in json_data['extractions']:
                    # print(keyword['parsed_value'])

    print("%d Document(s)" % len(f))
    # print(f[0])


if __name__ == '__main__':
    main()
